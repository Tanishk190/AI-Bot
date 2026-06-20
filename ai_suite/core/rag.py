"""RAG pipeline for document retrieval and chunking."""
from dataclasses import dataclass
from collections import Counter
import hashlib
import json
import math
import os
import re
from openai import OpenAI, APIError, APIConnectionError

try:
    from core.database import (
        get_embedding_by_hash, update_chunk_embedding, get_chunks_by_session,
    )
except ModuleNotFoundError:
    from ai_suite.core.database import (
        get_embedding_by_hash, update_chunk_embedding, get_chunks_by_session,
    )


@dataclass
class Chunk:
    text: str
    source: str
    index: int
    page: int | None = None


_embeddings_client = None
_embedding_model_name = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")


def _get_embeddings_client():
    global _embeddings_client
    if _embeddings_client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("Missing OPENAI_API_KEY.")
        _embeddings_client = OpenAI(api_key=api_key)
    return _embeddings_client


def _text_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _get_embedding(text: str) -> list[float]:
    h = _text_hash(text)
    cached = get_embedding_by_hash(h)
    if cached is not None:
        return cached

    try:
        client = _get_embeddings_client()
        response = client.embeddings.create(model=_embedding_model_name, input=text)
        embedding = response.data[0].embedding
        update_chunk_embedding(h, embedding)
        return embedding
    except APIConnectionError as exc:
        raise RuntimeError(f"Could not connect to OpenAI embeddings API: {str(exc)}") from exc
    except APIError as exc:
        raise RuntimeError(f"OpenAI embeddings API error: {str(exc)}") from exc


def embed_session_chunks(session_db_id: int):
    """Compute and store embeddings for all chunks in a session that don't have one yet."""
    chunk_rows = get_chunks_by_session(session_db_id)
    for row in chunk_rows:
        if row["embedding"] is None:
            embedding = _get_embedding(row["text"])
            # _get_embedding may return a cached value without updating this specific chunk
            update_chunk_embedding(row["text_hash"], embedding)


def _tokenize(text: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", text.lower())


def _keyword_scores(question: str, chunk_rows: list[dict]) -> list[float]:
    """BM25 keyword score for every chunk against the query.

    Gives proper weight to rare, high-signal terms (e.g. "269ss", "cheque")
    so keyword matches surface relevant chunks the embedding ranking may miss.
    """
    q_terms = set(_tokenize(question))
    docs = [_tokenize(r["text"]) for r in chunk_rows]
    n = len(docs)
    if n == 0 or not q_terms:
        return [0.0] * n

    avgdl = sum(len(d) for d in docs) / n or 1.0
    df: dict[str, int] = {}
    for d in docs:
        for term in set(d):
            df[term] = df.get(term, 0) + 1

    k1, b = 1.5, 0.75
    scores = []
    for d in docs:
        tf = Counter(d)
        dl = len(d) or 1
        score = 0.0
        for term in q_terms:
            freq = tf.get(term, 0)
            if freq:
                idf = math.log(1 + (n - df[term] + 0.5) / (df[term] + 0.5))
                score += idf * (freq * (k1 + 1)) / (freq + k1 * (1 - b + b * dl / avgdl))
        scores.append(score)
    return scores


def _semantic_scores(question: str, chunk_rows: list[dict]) -> list[float] | None:
    """Cosine similarity per chunk, or None if embeddings are unavailable."""
    try:
        import numpy as np
        q_vec = np.array(_get_embedding(question))
        q_norm = np.linalg.norm(q_vec)
        scores, any_emb = [], False
        for row in chunk_rows:
            emb = row["embedding"]
            if emb is None:
                scores.append(0.0)
                continue
            if isinstance(emb, str):
                emb = json.loads(emb)
            c_vec = np.array(emb)
            scores.append(float(np.dot(q_vec, c_vec) / (q_norm * np.linalg.norm(c_vec))))
            any_emb = True
        return scores if any_emb else None
    except Exception as e:
        print(f"Semantic scoring failed: {e}. Using keyword-only ranking.")
        return None


def _normalize(values: list[float]) -> list[float]:
    lo, hi = min(values), max(values)
    if hi <= lo:
        return [0.0] * len(values)
    return [(v - lo) / (hi - lo) for v in values]


def _detect_page_intent(question: str) -> set[int] | None:
    """Return the set of page numbers a query explicitly asks about, or None.

    Catches phrasings like "page 5", "on page 12", "pages 3-4", "p. 7",
    "page no. 9". When a user names a page, semantic/BM25 matching against the
    literal phrase "page 5" is useless (it isn't similar to the page's actual
    content), so we route these to a hard metadata filter instead. Returns None
    when no explicit page reference is found, leaving normal hybrid retrieval
    untouched.
    """
    pages: set[int] = set()
    # Ranges first: "pages 3-4", "page 3 to 5"
    for lo, hi in re.findall(r"\bpages?\s*(?:no\.?\s*)?(\d+)\s*(?:-|–|to)\s*(\d+)", question, re.I):
        lo_i, hi_i = int(lo), int(hi)
        if lo_i <= hi_i and hi_i - lo_i <= 50:  # guard against absurd ranges
            pages.update(range(lo_i, hi_i + 1))
    # Singles: full word "page 5"/"pages 5", or abbreviations "p. 7" / "pg 9".
    # Abbreviations require a trailing period or are matched as a standalone
    # token, so the bare "p" inside words like "top", "step", "group" is ignored.
    for n in re.findall(r"\bpages?\s*(?:no\.?\s*)?(\d+)", question, re.I):
        pages.add(int(n))
    for n in re.findall(r"\b(?:p\.\s*|pg\.?\s*)(\d+)", question, re.I):
        pages.add(int(n))
    return pages or None


def retrieve_relevant_chunks_db(question: str, session_db_id: int, k: int = 8,
                                document_ids: list[int] | None = None,
                                per_page_cap: int = 3,
                                semantic_weight: float = 0.6) -> list[Chunk]:
    """Hybrid retrieval over DB-stored chunks.

    Combines normalized semantic similarity with a BM25 keyword score, then
    spreads the results across pages (``per_page_cap``) so multi-clause queries
    aren't dominated by a single dense page.

    If the question explicitly names a page (e.g. "what's on page 5"), retrieval
    is hard-filtered to that page (plus immediately adjacent pages, since clauses
    span page breaks) before ranking — literal page queries don't match well on
    content similarity alone.
    """
    chunk_rows = get_chunks_by_session(session_db_id, document_ids=document_ids)
    if not chunk_rows:
        return []

    # Page-intent routing: restrict the candidate pool to requested pages (±1).
    requested_pages = _detect_page_intent(question)
    if requested_pages:
        wanted = set(requested_pages)
        for p in requested_pages:
            wanted.update({p - 1, p + 1})  # adjacent pages catch boundary-spanning clauses
        page_filtered = [r for r in chunk_rows if r.get("page") in wanted]
        # Only narrow if the requested pages actually exist; otherwise fall back
        # to the full pool so an out-of-range page number degrades gracefully.
        if page_filtered:
            chunk_rows = page_filtered

    keyword = _normalize(_keyword_scores(question, chunk_rows))
    semantic = _semantic_scores(question, chunk_rows)
    if semantic is not None:
        semantic = _normalize(semantic)
        combined = [semantic_weight * s + (1 - semantic_weight) * w
                    for s, w in zip(semantic, keyword)]
    else:
        combined = keyword

    order = sorted(range(len(chunk_rows)), key=lambda i: combined[i], reverse=True)

    # Page diversity: take the best chunks but cap how many come from one page
    selected, per_page = [], {}
    for i in order:
        page = chunk_rows[i].get("page")
        if per_page.get(page, 0) >= per_page_cap:
            continue
        selected.append(i)
        per_page[page] = per_page.get(page, 0) + 1
        if len(selected) >= k:
            break
    # Backfill if the per-page cap left us short of k
    if len(selected) < k:
        for i in order:
            if i not in selected:
                selected.append(i)
                if len(selected) >= k:
                    break

    return [
        Chunk(text=chunk_rows[i]["text"], source=chunk_rows[i]["source"],
              index=chunk_rows[i]["chunk_index"], page=chunk_rows[i].get("page"))
        for i in selected
    ]


# ── Document loading ─────────────────────────────────────────────────────────

def load_documents(uploaded_files: list) -> list[Chunk]:
    chunks = []
    for file_storage in uploaded_files:
        if not file_storage or not file_storage.filename:
            continue
        filename = os.path.basename(file_storage.filename)
        text = _extract_text(file_storage, filename)
        if text:
            chunks.extend(chunk_text(text, filename))
    return chunks


def load_chat_documents(uploaded_files: list) -> list[Chunk]:
    """Load chat documents into chunks, tagging each chunk with its source page.

    PDFs are chunked page-by-page so every chunk carries an accurate page
    number for citations. TXT/DOCX have no page concept, so page stays None.
    Chunk indices are renumbered sequentially across the whole document.
    """
    chunks = []
    for file_storage in uploaded_files:
        if not file_storage or not file_storage.filename:
            continue
        filename = os.path.basename(file_storage.filename)
        file_chunks: list[Chunk] = []

        if filename.lower().endswith(".pdf"):
            for page_num, page_text in _extract_pdf_pages(file_storage):
                if not page_text or not page_text.strip():
                    continue
                for chunk in chunk_text_structured(page_text, filename):
                    chunk.page = page_num
                    file_chunks.append(chunk)
            # Fallback: extractors that don't expose pages still yield whole-doc text
            if not file_chunks:
                text = _extract_text(file_storage, filename)
                if text:
                    file_chunks = chunk_text_structured(text, filename)
        else:
            text = _extract_text(file_storage, filename)
            if text:
                file_chunks = chunk_text_structured(text, filename)

        for i, chunk in enumerate(file_chunks, start=1):
            chunk.index = i
        chunks.extend(file_chunks)
    return chunks


def chunk_text(text: str, source: str, chunk_size: int = 600, overlap: int = 110) -> list[Chunk]:
    if not text or not text.strip():
        return []

    text = text.strip()
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    if len(paragraphs) == 1:
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    if len(paragraphs) == 1:
        paragraphs = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]

    chunks = []
    current = ""
    idx = 1

    for para in paragraphs:
        if len(para) > chunk_size:
            if current:
                chunks.append(Chunk(source=source, index=idx, text=current.strip()))
                idx += 1
            for i in range(0, len(para), chunk_size - overlap):
                piece = para[i: i + chunk_size]
                if piece.strip():
                    chunks.append(Chunk(source=source, index=idx, text=piece.strip()))
                    idx += 1
            current = ""
        elif len(current) + len(para) <= chunk_size:
            current += " " + para
        else:
            if current:
                chunks.append(Chunk(source=source, index=idx, text=current.strip()))
                idx += 1
                current = current[-overlap:] + " " + para
            else:
                current = para

    if current:
        chunks.append(Chunk(source=source, index=idx, text=current.strip()))

    return chunks


def chunk_text_structured(
    text: str, source: str, chunk_size: int = 600, overlap: int = 110
) -> list[Chunk]:
    if not text or not text.strip():
        return []

    sections = _split_sections(text)
    chunks = []
    idx = 1

    for heading, body in sections:
        if not body:
            continue

        if _is_table_block(body):
            chunks.append(Chunk(source=source, index=idx, text=_format_chunk(heading, body)))
            idx += 1
            continue

        clauses = _split_clauses(body)
        current = ""

        for clause in clauses:
            clause = clause.strip()
            if not clause:
                continue

            if len(clause) > chunk_size:
                if current:
                    chunks.append(Chunk(source=source, index=idx, text=_format_chunk(heading, current)))
                    idx += 1
                    current = ""
                step = max(chunk_size - overlap, 1)
                for i in range(0, len(clause), step):
                    piece = clause[i: i + chunk_size].strip()
                    if piece:
                        chunks.append(Chunk(source=source, index=idx, text=_format_chunk(heading, piece)))
                        idx += 1
                continue

            if not current:
                current = clause
            elif len(current) + len(clause) + 1 <= chunk_size:
                current = f"{current} {clause}"
            else:
                chunks.append(Chunk(source=source, index=idx, text=_format_chunk(heading, current)))
                idx += 1
                overlap_text = current[-overlap:].strip()
                current = f"{overlap_text} {clause}".strip() if overlap_text else clause

        if current:
            chunks.append(Chunk(source=source, index=idx, text=_format_chunk(heading, current)))
            idx += 1

    return chunks


def _is_table_block(text: str) -> bool:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if len(lines) < 3:
        return False
    tab_lines = sum(1 for l in lines if "\t" in l)
    pipe_lines = sum(1 for l in lines if "|" in l)
    if tab_lines >= 2 or pipe_lines >= 2:
        return True
    date_pattern = re.compile(r"\d{1,2}-[A-Za-z]+-\d{4}")
    date_lines = sum(1 for l in lines if date_pattern.search(l))
    if date_lines >= 2:
        return True
    lengths = [len(l) for l in lines]
    avg = sum(lengths) / len(lengths)
    return avg < 120 and (max(lengths) - min(lengths)) < 100


def _format_chunk(heading: str | None, text: str) -> str:
    if heading:
        return f"Section: {heading}\n{text}".strip()
    return text.strip()


def _split_sections(text: str) -> list[tuple[str | None, str]]:
    lines = text.splitlines()
    sections = []
    current_heading = None
    current_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_lines.append("")
            continue

        if _is_heading(stripped):
            if current_lines and any(part.strip() for part in current_lines):
                sections.append((current_heading, "\n".join(current_lines).strip()))
            current_heading = stripped
            current_lines = []
            continue

        current_lines.append(stripped)

    if current_lines and any(part.strip() for part in current_lines):
        sections.append((current_heading, "\n".join(current_lines).strip()))

    if not sections:
        return [(None, text.strip())]

    return sections


def _is_heading(line: str) -> bool:
    if len(line) > 140:
        return False
    heading_patterns = [
        r"^(section|article)\s+[0-9ivxlcdm]+([.-]\d+)*\b.*",
        r"^(schedule|appendix|exhibit)\s+[a-z0-9ivxlcdm]+.*",
        r"^\d+(\.\d+)*\s+[A-Z].{3,}$",
        r"^[A-Z][A-Z0-9 \-]{4,}$",
        r"^[A-Z][A-Za-z0-9 ,\-]{0,80}:$",
    ]
    for pattern in heading_patterns:
        if re.match(pattern, line.strip(), re.IGNORECASE):
            return True
    return False


def _split_clauses(text: str) -> list[str]:
    lines = text.splitlines()
    clauses = []
    current = ""
    clause_marker = re.compile(
        r"^\s*(\([a-z0-9]+\)|[a-z0-9]+\)|[a-z]\.|[0-9]+\.|[0-9]+\.[0-9]+)\s+",
        re.IGNORECASE
    )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                clauses.append(current.strip())
                current = ""
            continue

        if clause_marker.match(stripped):
            if current:
                clauses.append(current.strip())
            current = stripped
        else:
            current = f"{current} {stripped}".strip() if current else stripped

    if current:
        clauses.append(current.strip())

    if not clauses:
        return [text.strip()]

    refined = []
    for clause in clauses:
        if ";" in clause:
            parts = [p.strip() for p in clause.split(";") if p.strip()]
            refined.extend(parts)
        else:
            refined.append(clause)

    return refined


# ── In-memory retrieval (used by classifier with ephemeral chunks) ────────────

def retrieve_relevant_chunks(question: str, chunks: list[Chunk], k: int = 3) -> list[Chunk]:
    if not chunks:
        return []

    try:
        import numpy as np
        question_embedding = _get_embedding(question)
        scored = []
        for chunk in chunks:
            chunk_embedding = _get_embedding(chunk.text)
            similarity = np.dot(question_embedding, chunk_embedding) / (
                np.linalg.norm(question_embedding) * np.linalg.norm(chunk_embedding)
            )
            query_words = set(question.lower().split())
            chunk_words = set(chunk.text.lower().split())
            keyword_boost = 0.1 if query_words & chunk_words else 0
            scored.append((chunk, similarity + keyword_boost))

        scored.sort(key=lambda x: x[1], reverse=True)
        return [chunk for chunk, _ in scored[:k]]
    except Exception as e:
        print(f"Semantic search failed: {e}. Falling back to BM25.")

    return _bm25_retrieval(question, chunks, k)


def _bm25_retrieval(question: str, chunks: list[Chunk], k: int) -> list[Chunk]:
    query_words = set(question.lower().split())
    scored = []
    for chunk in chunks:
        chunk_words = set(chunk.text.lower().split())
        overlap = len(query_words & chunk_words)
        if overlap > 0:
            scored.append((chunk, overlap))
    scored.sort(key=lambda x: x[1], reverse=True)
    return [chunk for chunk, _ in scored[:k]]


# ── Text extraction (unchanged) ──────────────────────────────────────────────

def _extract_text(file_storage, filename: str) -> str:
    try:
        if filename.lower().endswith(".txt"):
            return file_storage.read().decode("utf-8", errors="ignore")
        elif filename.lower().endswith(".pdf"):
            return _extract_pdf(file_storage)
        elif filename.lower().endswith(".docx"):
            return _extract_docx(file_storage)
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
    return ""


def _table_to_text(table: list[list]) -> str:
    """Render an extracted table so column structure survives chunking.

    A flat tab/space join destroys the row↔column association — once chunked,
    the LLM sees a wall of numbers and can't tell which value sits under which
    header, which is exactly why broad questions over wide tables (e.g. "how
    much TDS was not deposited?") miss a single populated cell.

    We emit BOTH representations:
      1. A markdown table (pipe-delimited) so the visual grid is preserved.
      2. One self-describing line per data row that pairs every header with its
         cell value ("Section: 194I | TDS not deposited: 24,000 | ..."). Even
         if a row is split from the header during chunking, each cell still
         carries its own column label, so broad queries can't lose the link.
    """
    # Normalize: strip cells, collapse internal newlines.
    norm = []
    for row in table:
        if row is None:
            continue
        cells = [(c or "").strip().replace("\n", " ") for c in row]
        if any(cells):  # skip fully-empty rows
            norm.append(cells)
    if not norm:
        return ""

    width = max(len(r) for r in norm)
    norm = [r + [""] * (width - len(r)) for r in norm]  # pad ragged rows

    header = norm[0]
    body = norm[1:]

    # 1. Markdown grid
    md_lines = ["| " + " | ".join(header) + " |",
                "| " + " | ".join(["---"] * width) + " |"]
    for r in body:
        md_lines.append("| " + " | ".join(r) + " |")
    markdown = "\n".join(md_lines)

    # 2. Self-describing row lines (only when we have real headers to pair with)
    row_lines = []
    has_headers = any(h.strip() for h in header)
    if has_headers and body:
        for r in body:
            pairs = [f"{header[c].strip()}: {r[c].strip()}"
                     for c in range(width)
                     if header[c].strip() and r[c].strip()]
            if pairs:
                row_lines.append("Row — " + " | ".join(pairs))

    parts = [markdown]
    if row_lines:
        parts.append("\n".join(row_lines))
    return "\n\n".join(parts)


def _extract_pdf_pages(file_storage) -> list[tuple[int, str]]:
    """Extract text per page as (page_number, text), tables included.

    Tries pdfplumber → PyPDF2 → pypdf, mirroring _extract_pdf, but preserves
    page boundaries so chunks can be attributed to a page.
    """
    try:
        import pdfplumber
        try:
            file_storage.stream.seek(0)
        except Exception:
            pass
        pages_out: list[tuple[int, str]] = []
        with pdfplumber.open(file_storage.stream) as pdf:
            for n, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text(x_tolerance=1, y_tolerance=1) or ""
                table_blocks = []
                for table in (page.extract_tables() or []):
                    block = _table_to_text(table)
                    if block:
                        table_blocks.append(block)
                if table_blocks:
                    table_text = "\n\n".join(table_blocks)
                    page_text = f"{page_text}\n\n{table_text}".strip() if page_text else table_text
                pages_out.append((n, page_text))
        if any(t.strip() for _, t in pages_out):
            return pages_out
    except ImportError:
        pass
    except Exception as e:
        print(f"Per-page extraction with pdfplumber failed: {e}")

    try:
        from PyPDF2 import PdfReader
        try:
            file_storage.stream.seek(0)
        except Exception:
            pass
        reader = PdfReader(file_storage)
        pages_out = [(n, (page.extract_text() or "")) for n, page in enumerate(reader.pages, start=1)]
        if any(t.strip() for _, t in pages_out):
            return pages_out
    except ImportError:
        pass
    except Exception as e:
        print(f"Per-page extraction with PyPDF2 failed: {e}")

    try:
        from pypdf import PdfReader
        try:
            file_storage.stream.seek(0)
        except Exception:
            pass
        try:
            reader = PdfReader(file_storage, strict=False)
        except TypeError:
            reader = PdfReader(file_storage)
        out = []
        for n, page in enumerate(reader.pages, start=1):
            try:
                extracted = page.extract_text(extraction_mode="layout")
            except TypeError:
                extracted = page.extract_text()
            out.append((n, extracted or ""))
        return out
    except Exception as e:
        print(f"Per-page extraction with pypdf failed: {e}")
        return []


def _extract_pdf(file_storage) -> str:
    text = ""
    try:
        import pdfplumber
        try:
            file_storage.stream.seek(0)
        except Exception:
            pass
        with pdfplumber.open(file_storage.stream) as pdf:
            page_texts = []
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=1, y_tolerance=1) or ""
                tables = page.extract_tables() or []
                table_blocks = []
                for table in tables:
                    block = _table_to_text(table)
                    if block:
                        table_blocks.append(block)
                if table_blocks:
                    table_text = "\n\n".join(table_blocks)
                    page_text = f"{page_text}\n\n{table_text}".strip() if page_text else table_text
                if page_text:
                    page_texts.append(page_text)
            text = "\n\n".join(page_texts).strip()
            if text:
                return text
    except ImportError:
        pass
    except Exception as e:
        print(f"Error extracting text with pdfplumber: {e}")

    try:
        from PyPDF2 import PdfReader
        try:
            file_storage.stream.seek(0)
        except Exception:
            pass
        pdf = PdfReader(file_storage)
        for page in pdf.pages:
            text += page.extract_text() or ""
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"Error extracting text with PyPDF2: {e}")

    try:
        from pypdf import PdfReader
        try:
            file_storage.stream.seek(0)
        except Exception:
            pass
        try:
            pdf = PdfReader(file_storage, strict=False)
        except TypeError:
            pdf = PdfReader(file_storage)
        text = ""
        for page in pdf.pages:
            try:
                extracted = page.extract_text(extraction_mode="layout")
            except TypeError:
                extracted = page.extract_text()
            text += extracted or ""
        return text
    except ImportError:
        return "[PDF extraction requires PyPDF2 or pypdf]"
    except Exception as e:
        print(f"Error extracting text with pypdf: {e}")
        return text


def _extract_docx(file_storage) -> str:
    try:
        from docx import Document
        doc = Document(file_storage)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            grid = []
            for row in table.rows:
                grid.append([cell.text for cell in row.cells])
            block = _table_to_text(grid)
            if block:
                parts.append(block)
        return "\n".join([p for p in parts if p.strip()])
    except ImportError:
        return "[DOCX extraction requires python-docx]"